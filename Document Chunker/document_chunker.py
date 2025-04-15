import argparse
import os
import concurrent.futures
from pathlib import Path
from typing import Tuple, Optional
from dataclasses import dataclass
import time
import mmap
import io

# Import libraries with error handling
try:
    from PyPDF2 import PdfReader, PdfWriter
except ImportError:
    print("PyPDF2 not found. PDF functionality will not be available.")
    PdfReader = PdfWriter = None

try:
    import docx
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("python-docx not found. Word document functionality will not be available.")
    Document = None


@dataclass
class ChunkResult:
    """Data class for storing chunk results."""
    chunk_path: str
    start_index: int
    end_index: int
    success: bool
    error: Optional[str] = None


class DocumentChunker:
    """Base class for document chunking."""
    
    def __init__(self, input_path: str, output_dir: str, chunk_size: int, max_workers: int = None):
        self.input_path = input_path
        self.output_dir = output_dir
        self.chunk_size = chunk_size
        self.input_basename = Path(input_path).stem
        self.max_workers = max_workers or min(32, os.cpu_count() + 4)
        os.makedirs(output_dir, exist_ok=True)
    
    def process(self) -> Tuple[bool, str]:
        """Process the document and chunk it. To be implemented by subclasses."""
        raise NotImplementedError("Subclasses must implement this method")


class PDFChunker(DocumentChunker):
    """PDF document chunker."""
    
    def __init__(self, input_path: str, output_dir: str, chunk_size: int, max_workers: int = None):
        super().__init__(input_path, output_dir, chunk_size, max_workers)
        if PdfReader is None or PdfWriter is None:
            raise ImportError("PyPDF2 is required for PDF processing")
        
        # Use memory mapping for large files
        self.file_obj = open(input_path, 'rb')
        try:
            self.mm = mmap.mmap(self.file_obj.fileno(), 0, access=mmap.ACCESS_READ)
            self.pdf = PdfReader(self.mm)
        except:
            # Fallback if memory mapping isn't available
            self.mm = None
            self.file_obj.seek(0)
            self.pdf = PdfReader(self.file_obj)
            
        self.total_pages = len(self.pdf.pages)
    
    def __del__(self):
        """Clean up resources."""
        if hasattr(self, 'mm') and self.mm:
            self.mm.close()
        if hasattr(self, 'file_obj') and self.file_obj:
            self.file_obj.close()
    
    def _create_chunk(self, chunk_index: int) -> ChunkResult:
        """Create a single PDF chunk."""
        start_page = chunk_index * self.chunk_size
        end_page = min((chunk_index + 1) * self.chunk_size, self.total_pages)
        
        try:
            # Create a new PDF writer
            output = PdfWriter()
            
            # Add pages to the output
            for page_num in range(start_page, end_page):
                output.add_page(self.pdf.pages[page_num])
            
            # Save the chunk
            output_filename = f"{self.input_basename}_chunk_{chunk_index+1}.pdf"
            output_path = os.path.join(self.output_dir, output_filename)
            
            # Use a memory buffer to avoid file IO bottlenecks
            with io.BytesIO() as in_memory:
                output.write(in_memory)
                in_memory.seek(0)
                with open(output_path, "wb") as output_file:
                    output_file.write(in_memory.getbuffer())
            
            return ChunkResult(
                chunk_path=output_path,
                start_index=start_page + 1,
                end_index=end_page,
                success=True
            )
        except Exception as e:
            return ChunkResult(
                chunk_path="",
                start_index=start_page + 1,
                end_index=end_page,
                success=False,
                error=str(e)
            )
    
    def process(self) -> Tuple[bool, str]:
        """Process the PDF document and break it into chunks."""
        try:
            num_chunks = (self.total_pages + self.chunk_size - 1) // self.chunk_size
            result_message = f"Breaking {self.input_path} ({self.total_pages} pages) into {num_chunks} chunks...\n\n"
            
            failures = 0
            with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                futures = [executor.submit(self._create_chunk, i) for i in range(num_chunks)]
                
                for future in concurrent.futures.as_completed(futures):
                    result = future.result()
                    if result.success:
                        result_message += f"Created {result.chunk_path} (pages {result.start_index}-{result.end_index})\n"
                    else:
                        failures += 1
                        result_message += f"Error creating chunk {result.start_index}-{result.end_index}: {result.error}\n"
            
            if failures > 0:
                result_message += f"\nWarning: {failures} chunks failed to process.\n"
                
            return failures == 0, result_message
            
        except Exception as e:
            return False, f"Error processing PDF: {str(e)}"


class DocxChunker(DocumentChunker):
    """Word document chunker."""
    
    def __init__(self, input_path: str, output_dir: str, chunk_size: int, max_workers: int = None):
        super().__init__(input_path, output_dir, chunk_size, max_workers)
        if Document is None:
            raise ImportError("python-docx is required for Word document processing")
        
        # Load the document
        self.doc = Document(input_path)
        self.total_paragraphs = len(self.doc.paragraphs)
    
    def _create_chunk(self, chunk_index: int) -> ChunkResult:
        """Create a single Word document chunk."""
        start_para = chunk_index * self.chunk_size
        end_para = min((chunk_index + 1) * self.chunk_size, self.total_paragraphs)
        
        try:
            # Create a new document
            new_doc = Document()
            
            # Copy paragraphs from the original document
            for i in range(start_para, end_para):
                original_para = self.doc.paragraphs[i]
                
                # Create a new paragraph in the destination document
                new_para = new_doc.add_paragraph()
                
                # Copy style if available
                if original_para.style:
                    try:
                        new_para.style = original_para.style.name
                    except:
                        pass
                
                # Copy text with basic formatting
                for run in original_para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
            
            # Save the chunk
            output_filename = f"{self.input_basename}_chunk_{chunk_index+1}.docx"
            output_path = os.path.join(self.output_dir, output_filename)
            new_doc.save(output_path)
            
            return ChunkResult(
                chunk_path=output_path,
                start_index=start_para + 1,
                end_index=end_para,
                success=True
            )
        except Exception as e:
            return ChunkResult(
                chunk_path="",
                start_index=start_para + 1,
                end_index=end_para,
                success=False,
                error=str(e)
            )
    
    def process(self) -> Tuple[bool, str]:
        """Process the Word document and break it into chunks."""
        try:
            num_chunks = (self.total_paragraphs + self.chunk_size - 1) // self.chunk_size
            result_message = f"Breaking {self.input_path} ({self.total_paragraphs} paragraphs) into {num_chunks} chunks...\n\n"
            
            # Word document processing is resource-intensive, use fewer threads
            max_workers = min(4, self.max_workers)
            failures = 0
            
            with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = [executor.submit(self._create_chunk, i) for i in range(num_chunks)]
                
                for future in concurrent.futures.as_completed(futures):
                    result = future.result()
                    if result.success:
                        result_message += f"Created {result.chunk_path} (paragraphs {result.start_index}-{result.end_index})\n"
                    else:
                        failures += 1
                        result_message += f"Error creating chunk {result.start_index}-{result.end_index}: {result.error}\n"
            
            if failures > 0:
                result_message += f"\nWarning: {failures} chunks failed to process.\n"
                
            return failures == 0, result_message
            
        except Exception as e:
            return False, f"Error processing Word document: {str(e)}"


def get_chunker_for_file(file_path: str, output_dir: str, chunk_size: int, max_workers: int = None) -> DocumentChunker:
    """Factory function to get the appropriate chunker based on file extension."""
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf' and PdfReader is not None:
        return PDFChunker(file_path, output_dir, chunk_size, max_workers)
    elif file_ext == '.docx' and Document is not None:
        return DocxChunker(file_path, output_dir, chunk_size, max_workers)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}. Supported types are: .pdf, .docx")


def chunk_document(input_path: str, output_dir: str, chunk_size: int, max_workers: int = None) -> Tuple[bool, str]:
    """Break a document into chunks of specified size."""
    try:
        start_time = time.time()
        chunker = get_chunker_for_file(input_path, output_dir, chunk_size, max_workers)
        success, message = chunker.process()
        
        # Add timing information
        elapsed_time = time.time() - start_time
        message += f"\nProcessing completed in {elapsed_time:.2f} seconds."
        
        return success, message
    
    except Exception as e:
        return False, f"Error processing document: {str(e)}"


def main():
    parser = argparse.ArgumentParser(description="Break documents (PDF, DOCX) into chunks")
    parser.add_argument("input_file", help="Path to the input document file (PDF or DOCX)")
    parser.add_argument("--output-dir", default="chunks", 
                        help="Directory to save the chunked documents (default: 'chunks')")
    parser.add_argument("--chunk-size", type=int, default=2, 
                        help="Number of pages/paragraphs per chunk (default: 2)")
    parser.add_argument("--max-workers", type=int, default=None,
                        help="Maximum number of worker threads (default: CPU count + 4)")
    
    args = parser.parse_args()
    
    try:
        success, message = chunk_document(
            args.input_file, 
            args.output_dir, 
            args.chunk_size,
            args.max_workers
        )
        print(message)
        if not success:
            exit(1)
    except Exception as e:
        print(f"Error: {str(e)}")
        exit(1)


if __name__ == "__main__":
    main() 